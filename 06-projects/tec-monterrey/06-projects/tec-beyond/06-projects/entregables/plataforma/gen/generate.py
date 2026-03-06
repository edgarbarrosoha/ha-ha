"""Generate Plataforma Tec Beyond .docx using python-docx."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

script_dir = os.path.dirname(os.path.abspath(__file__))
logo_path = os.path.join(script_dir, "logo.png")
out_path = os.path.join(script_dir, "..", "Plataforma-Tec-Beyond-Descripcion-v1.docx")

# ── COLORS ──
HA_BLUE = RGBColor(0x33, 0x56, 0xF6)
HA_DARK = RGBColor(0x23, 0x1F, 0x20)
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xCC, 0xCC, 0xCC)

SANS = "Plus Jakarta Sans"
SANS_MED = "Plus Jakarta Sans Medium"
SANS_SEMI = "Plus Jakarta Sans SemiBold"
MONO = "Roboto Mono"

doc = Document()

# ── PAGE SETUP ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True


# ── HELPERS ──
def set_run(run, font_name=SANS, size=Pt(10.5), color=HA_DARK, bold=False):
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = False
    return run

def add_body(text, space_after=Pt(8)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    return p

def add_bullet(text, level=0, last=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12) if last else Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    for run in p.runs:
        run.clear()
    set_run(p.add_run(text))
    return p

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    pf = p.paragraph_format
    set_run(p.add_run(" "), size=Pt(2), color=WHITE)
    return p

def section_sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def section_num(n):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(str(n).zfill(2)), font_name=MONO, size=Pt(9), color=GREY)
    return p

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_MED
        run.font.size = Pt(20)
        run.font.color.rgb = HA_BLUE
        run.font.bold = False
        run.font.italic = False
    # Blue bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="6" w:color="3356F6"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_SEMI
        run.font.size = Pt(14)
        run.font.color.rgb = HA_DARK
        run.font.bold = False
        run.font.italic = False
    return p

def add_h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_SEMI
        run.font.size = Pt(12)
        run.font.color.rgb = HA_DARK
        run.font.bold = False
        run.font.italic = False
    return p

def add_quote(text, author=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    # Shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)
    # Left border blue
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="3356F6"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    # Content
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4) if author else Pt(6)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    if author:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        set_run(p2.add_run(author), font_name=MONO, size=Pt(8), color=GREY)
    return tbl

def add_data_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(h), font_name=SANS_SEMI, size=Pt(10), color=WHITE)
        # Blue shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3356F6" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_run(p.add_run(val), size=Pt(10))
    # Remove vertical borders, keep horizontal light
    tbl_el = tbl._element
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl_el.insert(0, tblPr)
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)
    return tbl

def add_code(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)
    # No borders
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.text = ""
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run("\n")
        set_run(p.add_run(line), font_name=MONO, size=Pt(7.5), color=HA_DARK)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return tbl

def mono_line(text, size=Pt(8), color=GREY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), font_name=MONO, size=size, color=color)
    return p

# ══════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════

# Logo
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run()
run.add_picture(logo_path, width=Inches(1.1))

add_spacer(100)

# Label
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("ENTREGABLE"), font_name=MONO, size=Pt(9), color=GREY)

# Title with blue bottom border
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
p.paragraph_format.line_spacing = 1.3
set_run(p.add_run("Plataforma Tec Beyond: descripción de aplicación y especificación técnica"), font_name=SANS_MED, size=Pt(20), color=BLACK)
pPr = p._element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="8" w:color="3356F6"/></w:pBdr>')
pPr.append(pBdr)

add_spacer(16)

# Metadata box
tbl = doc.add_table(rows=2, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri in range(2):
    for ci in range(2):
        cell = tbl.cell(ri, ci)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(borders)
# Remove table borders
tblPr = tbl._element.find(qn('w:tblPr'))
borders_xml = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders_xml)

labels = [("PREPARADO PARA:", "JORGE BLANDO"), ("PREPARADO POR:", "EDGAR BARROSO\nHORIZONS ARCHITECTURE")]
for ri, (label, val) in enumerate(labels):
    cell_l = tbl.cell(ri, 0)
    cell_l.paragraphs[0].text = ""
    set_run(cell_l.paragraphs[0].add_run(label), font_name=MONO, size=Pt(8), color=GREY)
    cell_r = tbl.cell(ri, 1)
    cell_r.paragraphs[0].text = ""
    for i, line in enumerate(val.split("\n")):
        if i > 0:
            cell_r.paragraphs[0].add_run("\n")
        set_run(cell_r.paragraphs[0].add_run(line), font_name=MONO, size=Pt(8), color=HA_DARK)

add_spacer(80)

# Footer of cover - 3 column table
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
items = [("HORIZONSARCHITECTURE.AI", WD_ALIGN_PARAGRAPH.LEFT), ("FECHA: MARZO 2026", WD_ALIGN_PARAGRAPH.CENTER), ("VERSIÓN: 1.0", WD_ALIGN_PARAGRAPH.RIGHT)]
for ci, (txt, align) in enumerate(items):
    cell = tbl.cell(0, ci)
    cell.paragraphs[0].text = ""
    cell.paragraphs[0].alignment = align
    set_run(cell.paragraphs[0].add_run(txt), font_name=MONO, size=Pt(8), color=GREY)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
tblPr = tbl._element.find(qn('w:tblPr'))
borders_xml = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders_xml)

# ══════════════════════════════════════════════
# PAGE BREAK → BODY
# ══════════════════════════════════════════════
doc.add_page_break()

# ── RUNNING HEADER (first page header is empty for cover) ──
section = doc.sections[0]
# Default header (pages after first)
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hp.text = ""
r1 = hp.add_run("2026 HORIZONS ARCHITECTURE")
r1.font.name = MONO
r1.font.size = Pt(8)
r1.font.color.rgb = GREY
# Add tab stop for right alignment
hp.add_run("\t\t")
r2 = hp.add_run("HORIZONSARCHITECTURE.AI")
r2.font.name = MONO
r2.font.size = Pt(8)
r2.font.color.rgb = GREY
# Set tab stops
pPr = hp._element.get_or_add_pPr()
tabs = parse_xml(
    f'<w:tabs {nsdecls("w")}>'
    '<w:tab w:val="center" w:pos="4680"/>'
    '<w:tab w:val="right" w:pos="9360"/>'
    '</w:tabs>'
)
pPr.append(tabs)

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.text = ""
r1 = fp.add_run("PLATAFORMA TEC BEYOND")
r1.font.name = MONO
r1.font.size = Pt(7)
r1.font.color.rgb = GREY
fp.add_run("\t\t")
r2 = fp.add_run("MARZO 2026")
r2.font.name = MONO
r2.font.size = Pt(7)
r2.font.color.rgb = GREY
pPr = fp._element.get_or_add_pPr()
tabs = parse_xml(
    f'<w:tabs {nsdecls("w")}>'
    '<w:tab w:val="center" w:pos="4680"/>'
    '<w:tab w:val="right" w:pos="9360"/>'
    '</w:tabs>'
)
pPr.append(tabs)

# First page header/footer = empty (for cover)
first_header = section.first_page_header
first_header.is_linked_to_previous = False
first_footer = section.first_page_footer
first_footer.is_linked_to_previous = False

# ══════════════════════════════════════════════
# BODY CONTENT
# ══════════════════════════════════════════════

# 01
section_num("01")
add_h1("Qué es")
add_body("Un espacio digital donde la comunidad Tec Beyond comparte conocimiento, se conecta por afinidad genuina y construye legado colectivo — no una red comercial, no un directorio, no un canal de ventas.")
add_spacer(6)
add_quote(
    "\"Nosotros venimos a esta comunidad no para vender los productos.\"\n\"El conocimiento compartido genera más conocimiento.\"",
    "— Consejo de representantes, 26 de febrero 2026"
)
add_spacer(10)
add_body("La plataforma tiene dos principios de diseño:")
add_bullet("La tecnología va a la persona usuaria, no al revés. El canal principal es WhatsApp — donde la comunidad ya vive. La aplicación web existe para quien quiera profundizar, pero ninguna interacción esencial requiere abrir una plataforma nueva.")
add_bullet("Las ventas son consecuencia, nunca causa. Todo lo que el sistema sugiere, muestra o conecta está orientado a aprendizaje, perspectiva y complementariedad — no a transacciones.", last=True)

# 02
section_sep()
section_num("02")
add_h1("Principios del consejo reflejados en la plataforma")
add_body("Estos principios, expresados por el consejo de representantes, son los criterios de diseño contra los que se evalúa cada funcionalidad:")
add_spacer(6)
add_data_table(
    ["Principio", "Implicación en la plataforma"],
    [
        ["Compartir, no vender", "No hay catálogos, no hay vitrinas de productos. Los perfiles describen conocimiento, experiencia y lo que cada quien busca aprender"],
        ["Valor desde el primer momento", "Al entrar, cada integrante ya encuentra contenido relevante, contexto de su red y conexiones sugeridas — sin esperas, sin promesas a futuro"],
        ["Anti-comercial", "No hay espacios publicitarios, no hay rankings de actividad comercial. Si alguien usa la plataforma para vender, el sistema lo detecta y el Consejo interviene"],
        ["Privacidad como confianza", "La información de cada integrante es suya. Cada persona decide qué comparte, con quién, y puede retirarlo. Los datos que alimentan a los agentes están anonimizados"],
        ["Sin encuestas, sin fricción", "La plataforma aprende de las interacciones naturales — no pide que se llenen formularios ni se respondan encuestas"],
        ["Aprendizaje para toda la vida", "No es solo networking empresarial. Es contexto: qué está pasando geopolíticamente, qué tendencias importan para el futuro de sus familias y comunidades"],
        ["Plataforma viva", "Los algoritmos evolucionan, las necesidades se mueven. La plataforma se actualiza continuamente — no es un producto que se entrega una vez"],
        ["Legado", "Lo que se comparte aquí cambia la vida de otros. La plataforma captura y preserva el conocimiento de la comunidad a través de generaciones"],
    ]
)

# 03
section_sep()
section_num("03")
add_h1("Qué ve cada persona")

add_h2("Integrante")
add_body("Al entrar por primera vez, cada integrante encuentra valor inmediato — como entrar a una sala donde ya hay algo esperando.")

add_h3("Inicio — La sala")
add_bullet("Contexto relevante desde el primer momento: qué está moviendo a la comunidad esta semana, qué temas conversan integrantes con perfiles similares, qué contenido se publicó recientemente en sus áreas de interés")
add_bullet("Conexiones sugeridas con explicación de por qué: \"Comparten experiencia en manufactura del Bajío y perspectivas complementarias sobre talento\"")
add_bullet("Sin métricas de actividad visibles, sin gamificación, sin presión para \"participar más\"", last=True)

add_h3("Mi perfil")
add_bullet("Contexto profesional: empresa, sector, trayectoria, áreas de experiencia")
add_bullet("Lo que sabe y puede compartir: conocimiento, perspectiva, red de contactos, mentoría")
add_bullet("Lo que quiere aprender o explorar: tendencias, mercados, capacidades, conexiones con otros contextos")
add_bullet("Control total de privacidad: cada integrante decide qué es visible para la red, qué es visible solo para el sistema (anonimizado), y qué no comparte")
add_bullet("Historial en Tec Beyond: generación, sede, participaciones", last=True)

add_h3("Conexiones")
add_bullet("Sugerencias del agente basadas en complementariedad de conocimiento y contexto — no en potencial de negocio")
add_bullet("Cada sugerencia explica la razón: \"Dos personas que lideran empresas familiares en transición generacional y podrían compartir perspectivas\"")
add_bullet("Double opt-in: ninguna parte ve quién es la otra hasta que ambas aceptan")
add_bullet("Se rechaza sin consecuencias — la otra persona nunca se entera")
add_bullet("Historial de conexiones anteriores", last=True)

add_h3("Contenido y contexto")
add_bullet("Artículos, análisis, casos y recursos seleccionados según su perfil e intereses — no genéricos")
add_bullet("Contenido diferenciado: una persona con 3 años de experiencia recibe material distinto de quien lleva 15")
add_bullet("Contexto vivo: qué está pasando en el entorno de negocios de su región, qué temas emergen en la comunidad")
add_bullet("Acceso a materiales de sesiones pasadas de Tec Beyond", last=True)

add_h3("Directorio")
add_bullet("Búsqueda de integrantes por sede, generación, industria, área de conocimiento")
add_bullet("Cada integrante configura su nivel de visibilidad", last=True)

add_h3("Eventos")
add_bullet("Calendario de encuentros de la comunidad (presenciales y virtuales)")
add_bullet("Registro directo")
add_bullet("Historial de participación", last=True)

add_h3("Vía WhatsApp — canal principal")
add_bullet("Todas las interacciones esenciales funcionan por WhatsApp: recibir sugerencias de conexión, aceptar o rechazar, recibir contenido relevante, confirmar asistencia a eventos")
add_bullet("La persona responde directamente en su conversación — sin descargar apps, sin crear cuentas nuevas, sin aprender una interfaz")
add_bullet("La aplicación web complementa para quien quiera explorar a fondo, pero no es requisito", last=True)

add_h2("Administración de sede")
add_h3("Dashboard local")
add_bullet("Salud de su comunidad: integrantes con actividad, tendencias de participación, temas emergentes")
add_bullet("Conexiones facilitadas en su sede este periodo")
add_bullet("Eventos locales y asistencia")
add_bullet("Alertas de integrantes que se desconectan (para acercamiento personal, no automatizado)", last=True)

add_h3("Gestión de integrantes")
add_bullet("Altas y bajas")
add_bullet("Seguimiento de onboarding (sin presionar — el sistema detecta si alguien necesita ayuda para empezar)", last=True)
add_spacer(4)
add_quote("Solo ve datos de su sede. No tiene acceso a integrantes de otras sedes ni a datos del Consejo.")

add_h2("El Consejo")
add_h3("Dashboard estratégico")
add_bullet("Visión de la red completa: integrantes por sede, generación, sector, áreas de conocimiento")
add_bullet("Salud de la comunidad: participación, conexiones realizadas, retención")
add_bullet("Tendencias detectadas: \"47 integrantes mencionaron nearshoring este trimestre\"")
add_bullet("Lo que la comunidad quiere aprender: temas emergentes, preguntas frecuentes, vacíos de conocimiento", last=True)

add_h3("Insights del agente")
add_bullet("Patrones en la red: sectores emergentes, oportunidades de conexión cross-sede, áreas donde la comunidad necesita más diversidad")
add_bullet("Reportes periódicos generados automáticamente — sin requerir encuestas")
add_bullet("Alertas: \"La sede Chihuahua tiene 30% menos participación que el promedio\"", last=True)

add_h3("Pipeline de candidaturas")
add_bullet("Personas candidatas sugeridas por el Agente Scout a partir de las redes de integrantes actuales")
add_bullet("Perfil estimado de cada persona candidata (trayectoria, sector, afinidad con la comunidad)")
add_bullet("El Consejo decide a quién invitar — el agente solo sugiere, nunca contacta", last=True)

add_h3("Gobernanza")
add_bullet("Reglas de la comunidad: criterios de admisión, políticas de conexión, niveles de membresía")
add_bullet("Registro de decisiones y acuerdos")
add_bullet("Auditoría: quién accedió a qué, cuándo, con qué resultado", last=True)

add_h3("Legado de la comunidad")
add_bullet("Conocimiento acumulado a lo largo de las generaciones: qué temas se han discutido, qué conexiones se han formado, qué impacto ha tenido la red")
add_bullet("Memoria institucional que trasciende a quienes integran la red — lo que alguien de generación 1 compartió está disponible como contexto para generación 10", last=True)

# 04
section_sep()
section_num("04")
add_h1("Flujos principales")

add_h3("Flujo 1: Primer contacto — valor visible")
add_code("La persona recibe invitación por WhatsApp (canal natural)\n    │\nResponde en WhatsApp: acepta, completa perfil básico conversando con el bot\n    │\nBot: contexto profesional, qué sabe/puede compartir, qué quiere aprender\n    │\nConfigura privacidad: qué es visible, qué es solo para el sistema (anonimizado)\n    │\nDe inmediato recibe: un contenido relevante + su primera sugerencia de conexión\n    │\nSi quiere profundizar → accede a la aplicación web. Si no, todo funciona por WhatsApp.")

add_h3("Flujo 2: Conexión entre integrantes")
add_code("Agente detecta complementariedad de conocimiento entre dos integrantes\n    │\nA cada persona (por WhatsApp): \"Hay alguien en la comunidad que comparte\ntu interés en [tema]. Tienen perspectivas complementarias. ¿Te gustaría conocerle?\"\n    │\n    ├── Ambas partes aceptan → Se facilita la introducción\n    ├── Una parte rechaza → No pasa nada. La otra persona nunca se entera.\n    └── Nadie responde → Se archiva. Se puede sugerir de nuevo más adelante.")

add_h3("Flujo 3: Detección de personas candidatas")
add_code("Agente Scout analiza patrones en la red de integrantes actuales\n    │\nIdentifica perfil con afinidad a la comunidad (trayectoria, valores, sector)\n    │\nPresenta al Consejo: \"En la red de [Integrante X] hay alguien con perfil afín\"\n    │\nConsejo decide si lo invita para la siguiente generación")

add_h3("Flujo 4: Inteligencia de la comunidad (sin encuestas)")
add_code("Los agentes analizan datos agregados y anonimizados de toda la red\n    │\nDetectan patrón: \"El 60% de integrantes de gen. 7-9 quieren aprender\nsobre IA aplicada a manufactura\"\n    │\nGeneran reporte para Consejo: \"Considerar evento especializado en IA+manufactura\"\n    │\nConsejo decide si actúa — la comunidad nunca fue interrumpida con una encuesta")

# 05
section_sep()
section_num("05")
add_h1("Métricas de éxito del piloto (León)")
add_spacer(6)
add_data_table(
    ["Indicador", "Meta", "Cómo se mide"],
    [
        ["Perfiles completos", "80% de integrantes", "Campos esenciales completados (vía WhatsApp o web)"],
        ["Integrantes que encuentran valor", "70% interactúa al menos 1 vez/mes", "Interacción por WhatsApp o web — sin exigir login"],
        ["Conexiones facilitadas", "50+ en 6 meses", "Conexiones aceptadas por ambas partes"],
        ["Percepción de valor", "NPS > 40", "Conversación breve del bot (no encuesta formal)"],
        ["Continuidad", "90% permanece al año", "Participación activa sostenida"],
    ]
)

# 06
section_sep()
section_num("06")
add_h1("Arquitectura general")
add_body("El sistema se organiza en tres capas: infraestructura del Tec (Azure), capa de agentes operada por HA, y las interfaces de usuario.")
add_spacer(6)
add_code("INFRAESTRUCTURA TEC (Azure)\n┌──────────────────────────────────────────────────────┐\n│  PostgreSQL      Vector DB        Blob Storage        │\n│  Perfiles        (Pinecone/       Documentos          │\n│  Membresías      pgvector)        Reportes            │\n│  Conexiones      Embeddings       Backups             │\n│  Eventos         Relaciones semánticas                │\n│  Auditoría                                            │\n└──────────────────────────────────────────────────────┘\n                        │\n                  API REST Gateway\n                  OAuth 2.0 / JWT\n                        │\n┌──────────────────────────────────────────────────────┐\n│              CAPA AGENTES (operada por HA)            │\n│   Conexión  │  Contenidos  │  Scout  │  Insights     │\n│  NO almacena datos de integrantes                     │\n└──────────────────────────────────────────────────────┘\n         │              │              │\n    Web App        WhatsApp Bot    Dashboard Consejo\n    (React)        (Twilio/360)    (Metabase/custom)")

# 07
section_sep()
section_num("07")
add_h1("Stack tecnológico")
add_spacer(6)
add_data_table(
    ["Capa", "Tecnología", "Justificación"],
    [
        ["Backend", "Node.js / Python (FastAPI)", "Ecosistema maduro, compatible con Azure"],
        ["BD relacional", "PostgreSQL (Azure Database)", "Estándar del Tec, robusto, escalable"],
        ["BD vectorial", "pgvector o Pinecone", "Búsqueda semántica para matching y contenido"],
        ["LLM", "Claude / GPT-5.2", "Procesamiento de lenguaje natural para agentes"],
        ["Orquestación agentes", "HA framework", "Coordinación multi-agente con reglas de negocio"],
        ["Frontend web", "React + Tailwind CSS", "Responsive, rendimiento, accesibilidad"],
        ["WhatsApp", "Business API (Twilio/360dialog)", "Canal preferido de la comunidad"],
        ["Dashboard", "Metabase o custom", "Visualización de insights para Consejo"],
        ["Autenticación", "OAuth 2.0 + JWT", "SSO compatible con ecosistema Tec (Azure AD)"],
        ["Infraestructura", "Microsoft Azure", "Alineado con infraestructura existente del Tec"],
        ["CI/CD", "GitHub Actions → Azure App Service", "Despliegue automatizado"],
        ["Monitoreo", "Azure Monitor + App Insights", "Observabilidad y alertas"],
    ]
)

# 08
section_sep()
section_num("08")
add_h1("Modelo de datos (entidades principales)")
add_spacer(6)
add_code("INTEGRANTE\n├── id, nombre, email, teléfono\n├── empresa, cargo, sector, ubicación\n├── sede_id, generación, año_ingreso\n├── conocimiento_experiencia (texto libre → embedding)\n├── que_quiere_aprender (texto libre → embedding)\n├── preferencias_privacidad (granular: visible / solo sistema / privado)\n├── canal_preferido (whatsapp / web / ambos)\n├── estado (activo / inactivo / pendiente)\n└── fecha_registro, última_actividad\n\nSEDE\n├── id, nombre, campus, ciudad\n├── admin_id, fecha_inicio\n└── estado (activa / en preparación)\n\nCONEXIÓN\n├── id, integrante_a_id, integrante_b_id\n├── score_afinidad (0-100)\n├── base_conexión (texto: por qué se sugiere)\n├── estado_a, estado_b (pendiente / aceptado / rechazado)\n├── estado_final (facilitado / cancelado / archivado)\n├── facilitado_por (admin_id)\n└── fechas (creado, respondido_a, respondido_b, facilitado)\n\nCONTENIDO\n├── id, título, tipo (artículo / caso / recurso / video)\n├── sectores_relevantes [], nivel_experiencia\n├── embedding (para matching con perfiles)\n└── fecha_publicación, fuente\n\nCANDIDATURA_SCOUT\n├── id, nombre_estimado, empresa, sector\n├── detectado_vía_integrante_id\n├── score_afinidad (0-100)\n├── estado (detectado / revisado / invitado / descartado)\n└── decidido_por (consejo), fecha_decisión\n\nEVENTO\n├── id, nombre, tipo (presencial / virtual), sede_id\n├── fecha, lugar, capacidad\n└── asistentes [] (integrante_id + confirmación)\n\nLOG_AUDITORÍA\n├── id, persona_id, acción, recurso\n├── timestamp, ip, resultado\n└── detalle (JSON)")

# 09
section_sep()
section_num("09")
add_h1("Integraciones")
add_spacer(6)
add_data_table(
    ["Sistema", "Tipo", "Qué se intercambia", "Dirección"],
    [
        ["Azure AD (Tec)", "SSO", "Autenticación de cuentas institucionales", "Tec → Plataforma"],
        ["WhatsApp Business API", "Mensajería", "Notificaciones, respuestas, flujos conversacionales", "Bidireccional"],
        ["LinkedIn (opcional)", "Enriquecimiento", "Datos públicos de perfil profesional", "LinkedIn → Plataforma"],
        ["Anthropic / OpenAI API", "LLM", "Procesamiento de lenguaje para agentes", "Plataforma → API"],
        ["SMTP / SendGrid", "Email", "Notificaciones, reportes, invitaciones", "Plataforma → Integrante"],
        ["Azure Blob Storage", "Archivos", "Documentos, reportes, backups", "Interno"],
    ]
)

# 10
section_sep()
section_num("10")
add_h1("Seguridad y confianza")
add_spacer(4)
add_quote(
    "\"Mucha gente no respondió la encuesta cuando leyó los avisos de privacidad del Tec. Tiene una vulnerabilidad reconocida en la comunidad.\"",
    "— Rodrigo, Consejo de representantes"
)
add_spacer(10)
add_body("La comunidad ya expresó desconfianza hacia el manejo de datos del Tec. La plataforma debe ganarse esa confianza con hechos, no con avisos legales.")
add_spacer(4)
add_quote("Principio: Cada integrante controla su información. Decide qué comparte, con quién, y puede retirarlo en cualquier momento. Los agentes de IA trabajan con datos anonimizados — nunca ven nombres, emails ni datos de contacto.")
add_spacer(10)
add_data_table(
    ["Capa", "Medida", "Estándar"],
    [
        ["Datos en reposo", "AES-256", "Azure Storage Service Encryption"],
        ["Datos en tránsito", "TLS 1.3", "HTTPS obligatorio"],
        ["Anonimización para IA", "Datos procesados sin PII antes de enviar a LLM", "Política interna verificable"],
        ["Autenticación admin", "MFA", "Azure AD + Authenticator"],
        ["API", "JWT con rotación de tokens", "Expiración configurable"],
        ["Acceso a datos", "RBAC", "6 roles definidos"],
        ["Logs", "Registro inmutable de toda operación", "Azure Monitor + tabla interna"],
        ["Backups", "Automático diario, retención 30 días", "Azure Backup"],
        ["Pruebas de seguridad", "Pentest antes de lanzamiento", "Proveedor externo"],
        ["Transparencia", "Cada integrante puede ver y exportar su información", "Panel de privacidad en perfil"],
    ]
)

add_h2("Roles y permisos")
add_spacer(6)
add_data_table(
    ["Rol", "Datos integrantes", "Conexiones", "Insights", "Config", "Auditoría"],
    [
        ["Super Admin", "Todos", "Todas", "Todos", "Sí", "Sí"],
        ["Admin Consejo", "Agregados", "Aprobar", "Todos", "Reglas", "Sí"],
        ["Admin Sede", "Solo su sede", "Solo su sede", "Solo su sede", "No", "Solo su sede"],
        ["Operación HA", "No", "No", "No", "Técnica", "Técnica"],
        ["Integrante", "Solo propio", "Propias", "No", "Su perfil", "No"],
        ["Auditoría", "Lectura", "Lectura", "Lectura", "No", "Sí"],
    ]
)

# 11
section_sep()
section_num("11")
add_h1("Residencia y propiedad de datos")
add_spacer(6)
add_data_table(
    ["Componente", "Ubicación", "Propiedad", "Si termina el contrato"],
    [
        ["Datos de integrantes", "Azure Tec", "Tec Beyond 100%", "Se revoca acceso a HA; datos intactos"],
        ["Base vectorial", "Azure Tec", "Tec Beyond 100%", "Se entrega documentación para migración"],
        ["Insights y reportes", "Azure Tec", "Tec Beyond 100%", "Quedan en infraestructura Tec"],
        ["Código de plataforma", "Repositorio HA", "Horizons Architecture", "Licencia se termina"],
        ["Algoritmos de agentes", "Repositorio HA", "Horizons Architecture", "HA retira su código; resultados quedan"],
        ["Reglas de comunidad", "En plataforma", "Consejo León", "Exportables en formato estándar"],
    ]
)
add_spacer(10)
add_quote("Compromiso de portabilidad: Si Tec Beyond migra a otro proveedor, HA entrega todos los datos en formato estándar (CSV/JSON), documentación técnica de la base vectorial, y un periodo de transición acordado en contrato.")

# 12
section_sep()
section_num("12")
add_h1("Escalabilidad")

add_h3("Piloto (León)")
add_data_table(
    ["Recurso", "Dimensionamiento"],
    [
        ["Conexiones simultáneas", "~50-100"],
        ["Integrantes totales", "~300"],
        ["Embeddings", "~300 perfiles + contenido"],
        ["Almacenamiento", "< 10 GB"],
        ["Costo mensual estimado Azure", "~$200-400 USD"],
        ["Costo LLM mensual estimado", "~$100-300 USD"],
    ]
)

add_h3("Expansión (3-5 sedes)")
add_data_table(
    ["Recurso", "Dimensionamiento"],
    [
        ["Conexiones simultáneas", "~200-500"],
        ["Integrantes totales", "~1,500"],
        ["Costo marginal por sede nueva", "~20-30% del costo del piloto"],
    ]
)
add_spacer(6)
add_body("La arquitectura es horizontal: agregar una sede es crear un nuevo nodo con su admin, no redesplegar el sistema. La base de conocimiento es compartida — cada nueva persona enriquece las conexiones de toda la red.")

# 13
section_sep()
section_num("13")
add_h1("Fases de desarrollo")
add_spacer(4)
add_quote(
    "\"Queremos que no nos deje el Tec con una plataforma una vez. Yo quiero quedarme con alguien que piensa generar.\"",
    "— Mario Orozco"
)
add_spacer(10)
add_data_table(
    ["Fase", "Duración", "Entregables"],
    [
        ["Descubrimiento", "2-3 semanas", "Diseño detallado, modelo de datos final, cotización cerrada"],
        ["MVP", "8-10 semanas", "Perfiles, conexiones básicas, WhatsApp bot, dashboard Consejo"],
        ["Agentes v1", "4-6 semanas", "Conexiones + Contenidos operando con base de conocimiento real"],
        ["Agentes v2", "4-6 semanas", "Scout + Insights operando"],
        ["Estabilización", "2-4 semanas", "Testing, pentest, ajustes de performance, onboarding"],
        ["Lanzamiento piloto", "—", "~300 integrantes de León participando"],
    ]
)
add_spacer(6)
add_quote("Total estimado: 5-8 meses desde aprobación hasta lanzamiento del piloto.")

add_h3("Después del lanzamiento: plataforma viva")
add_body("La plataforma no se \"entrega\" — se opera y evoluciona continuamente:")
add_bullet("Actualizaciones mensuales de algoritmos de conexión y contenido basadas en patrones reales de uso")
add_bullet("Reportes trimestrales al Consejo con insights y recomendaciones de evolución")
add_bullet("Nuevas funcionalidades según las necesidades que emerjan de la comunidad")
add_bullet("Expansión a nuevas sedes como nodos adicionales — cada sede nueva enriquece la red completa", last=True)

# 14
section_sep()
section_num("14")
add_h1("Preguntas frecuentes para TI")

add_h3("¿Se necesita infraestructura nueva en el Tec?")
add_body("No. Todo corre sobre Azure, que el Tec ya tiene. Se crea un resource group dedicado para Tec Beyond.")

add_h3("¿Los agentes de IA corren dentro de Azure?")
add_body("Los agentes se despliegan en Azure App Service dentro del tenant del Tec. Las llamadas a LLMs salen a APIs externas, pero los datos que se envían son procesados y anonimizados — nunca se envía un perfil completo a un LLM externo.")

add_h3("¿Qué pasa con la latencia de WhatsApp?")
add_body("El bot responde en < 5 segundos para interacciones simples (aceptar conexión, confirmar evento). Para respuestas que requieren procesamiento de agentes, < 15 segundos.")

add_h3("¿Cómo se actualiza la base vectorial?")
add_body("Los embeddings se regeneran cuando alguien actualiza su perfil. El sistema también re-indexa periódicamente (semanal) para incorporar datos de interacciones recientes.")

add_h3("¿Qué datos se envían a las APIs de LLM?")
add_body("Solo contexto procesado y anonimizado. Para generar una conexión, el agente envía al LLM las descripciones de capacidades y necesidades sin nombres, emails ni datos de contacto. El LLM nunca ve datos personales identificables.")

add_h3("¿Se puede integrar con sistemas existentes del Tec (CRM, ERP)?")
add_body("Sí, vía API REST. La fase de descubrimiento mapea qué integraciones específicas se necesitan.")

# Colophon
add_spacer(24)
p = doc.add_paragraph()
pPr = p._element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/></w:pBdr>')
pPr.append(pBdr)
add_spacer(10)
add_body("Documento preparado por Horizons Architecture para Tec Beyond.")
mono_line("Marzo 2026 — horizonsarchitecture.ai")

# ── SAVE ──
doc.save(out_path)
print(f"Generated: {out_path}")
