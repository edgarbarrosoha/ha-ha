"""Generate Especificación Técnica Tec Beyond .docx using python-docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

script_dir = os.path.dirname(os.path.abspath(__file__))
logo_path = os.path.join(script_dir, "logo.png")
out_path = os.path.join(script_dir, "..", "Plataforma-Tec-Beyond-Especificacion-Tecnica-v1.docx")

# ── COLORS ──
HA_BLUE = RGBColor(0x33, 0x56, 0xF6)
HA_DARK = RGBColor(0x23, 0x1F, 0x20)
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xCC, 0xCC, 0xCC)

SANS = "Plus Jakarta Sans"
SANS_MED = "Plus Jakarta Sans Medium"
SANS_SEMI = "Plus Jakarta Sans SemiBold"
MONO = "Roboto Mono"

doc = Document()

# ── PAGE SETUP ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True


# ── HELPERS ──
def set_run(run, font_name=SANS, size=Pt(10.5), color=HA_DARK, bold=False):
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = False
    return run

def add_body(text, space_after=Pt(8)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    return p

def add_body_bold_prefix(bold_text, normal_text, space_after=Pt(8)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(bold_text), bold=True)
    set_run(p.add_run(normal_text))
    return p

def add_bullet(text, level=0, last=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12) if last else Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    for run in p.runs:
        run.clear()
    set_run(p.add_run(text))
    return p

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(" "), size=Pt(2), color=WHITE)
    return p

def section_sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def section_num(n):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(str(n).zfill(2)), font_name=MONO, size=Pt(9), color=GREY)
    return p

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_MED
        run.font.size = Pt(20)
        run.font.color.rgb = HA_BLUE
        run.font.bold = False
        run.font.italic = False
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="6" w:color="3356F6"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_SEMI
        run.font.size = Pt(14)
        run.font.color.rgb = HA_DARK
        run.font.bold = False
        run.font.italic = False
    return p

def add_h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = SANS_SEMI
        run.font.size = Pt(12)
        run.font.color.rgb = HA_DARK
        run.font.bold = False
        run.font.italic = False
    return p

def add_quote(text, author=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="3356F6"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4) if author else Pt(6)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    if author:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        set_run(p2.add_run(author), font_name=MONO, size=Pt(8), color=GREY)
    return tbl

def add_data_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(h), font_name=SANS_SEMI, size=Pt(10), color=WHITE)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3356F6" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_run(p.add_run(val), size=Pt(10))
    tbl_el = tbl._element
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl_el.insert(0, tblPr)
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)
    return tbl

def add_code(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.text = ""
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run("\n")
        set_run(p.add_run(line), font_name=MONO, size=Pt(7.5), color=HA_DARK)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return tbl

def mono_line(text, size=Pt(8), color=GREY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), font_name=MONO, size=size, color=color)
    return p

def remove_table_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)
    tblPr = tbl._element.find(qn('w:tblPr'))
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


# ══════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════

# Logo
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run()
run.add_picture(logo_path, width=Inches(1.1))

add_spacer(100)

# Label
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("ENTREGABLE"), font_name=MONO, size=Pt(9), color=GREY)

# Title
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
p.paragraph_format.line_spacing = 1.3
set_run(p.add_run("Plataforma Tec Beyond: Especificación Técnica"), font_name=SANS_MED, size=Pt(20), color=BLACK)
pPr = p._element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="8" w:color="3356F6"/></w:pBdr>')
pPr.append(pBdr)

add_spacer(16)

# Metadata box
tbl = doc.add_table(rows=3, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri in range(3):
    for ci in range(2):
        cell = tbl.cell(ri, ci)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F2F2" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
remove_table_borders(tbl)

labels = [
    ("PARA:", "EQUIPO DE TI DEL TEC DE MONTERREY\nOSCAR DÍAZ (HORIZONS ARCHITECTURE)"),
    ("DE:", "EDGAR BARROSO\nHORIZONS ARCHITECTURE"),
    ("VERSIÓN:", "1.0"),
]
for ri, (label, val) in enumerate(labels):
    cell_l = tbl.cell(ri, 0)
    cell_l.paragraphs[0].text = ""
    set_run(cell_l.paragraphs[0].add_run(label), font_name=MONO, size=Pt(8), color=GREY)
    cell_r = tbl.cell(ri, 1)
    cell_r.paragraphs[0].text = ""
    for i, line in enumerate(val.split("\n")):
        if i > 0:
            cell_r.paragraphs[0].add_run("\n")
        set_run(cell_r.paragraphs[0].add_run(line), font_name=MONO, size=Pt(8), color=HA_DARK)

add_spacer(80)

# Footer of cover
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
items = [("HORIZONSARCHITECTURE.AI", WD_ALIGN_PARAGRAPH.LEFT), ("FECHA: MARZO 2026", WD_ALIGN_PARAGRAPH.CENTER), ("VERSIÓN: 1.0", WD_ALIGN_PARAGRAPH.RIGHT)]
for ci, (txt, align) in enumerate(items):
    cell = tbl.cell(0, ci)
    cell.paragraphs[0].text = ""
    cell.paragraphs[0].alignment = align
    set_run(cell.paragraphs[0].add_run(txt), font_name=MONO, size=Pt(8), color=GREY)
remove_table_borders(tbl)

# ══════════════════════════════════════════════
# PAGE BREAK → BODY
# ══════════════════════════════════════════════
doc.add_page_break()

# ── RUNNING HEADER ──
section = doc.sections[0]
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hp.text = ""
r1 = hp.add_run("2026 HORIZONS ARCHITECTURE")
r1.font.name = MONO
r1.font.size = Pt(8)
r1.font.color.rgb = GREY
hp.add_run("\t\t")
r2 = hp.add_run("HORIZONSARCHITECTURE.AI")
r2.font.name = MONO
r2.font.size = Pt(8)
r2.font.color.rgb = GREY
pPr = hp._element.get_or_add_pPr()
tabs = parse_xml(
    f'<w:tabs {nsdecls("w")}>'
    '<w:tab w:val="center" w:pos="4680"/>'
    '<w:tab w:val="right" w:pos="9360"/>'
    '</w:tabs>'
)
pPr.append(tabs)

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.text = ""
r1 = fp.add_run("ESPECIFICACIÓN TÉCNICA — TEC BEYOND")
r1.font.name = MONO
r1.font.size = Pt(7)
r1.font.color.rgb = GREY
fp.add_run("\t\t")
r2 = fp.add_run("MARZO 2026")
r2.font.name = MONO
r2.font.size = Pt(7)
r2.font.color.rgb = GREY
pPr = fp._element.get_or_add_pPr()
tabs = parse_xml(
    f'<w:tabs {nsdecls("w")}>'
    '<w:tab w:val="center" w:pos="4680"/>'
    '<w:tab w:val="right" w:pos="9360"/>'
    '</w:tabs>'
)
pPr.append(tabs)

# First page header/footer = empty (for cover)
first_header = section.first_page_header
first_header.is_linked_to_previous = False
first_footer = section.first_page_footer
first_footer.is_linked_to_previous = False

# ══════════════════════════════════════════════
# BODY CONTENT
# ══════════════════════════════════════════════

# ── 01 RESUMEN EJECUTIVO TÉCNICO ──
section_num("01")
add_h1("Resumen ejecutivo técnico")
add_body("Plataforma de comunidad inteligente para ~300 integrantes (piloto León), escalable a 3,000+ a nivel nacional. Opera en dos canales (WhatsApp + web). Cuatro agentes de IA procesan datos anonimizados para facilitar conexiones, personalizar contenido, detectar candidaturas y generar inteligencia para el Consejo. Toda la infraestructura corre sobre Azure del Tec. HA opera la capa de agentes; los datos son propiedad 100% del Tec.")
add_spacer(6)
add_quote("Este documento es el complemento técnico de la propuesta para liderazgo. Describe la arquitectura, los agentes, el modelo de datos, la seguridad y el plan de despliegue de la plataforma Tec Beyond.")

# ── 02 ARQUITECTURA GENERAL ──
section_sep()
section_num("02")
add_h1("Arquitectura general")
add_body("El sistema se organiza en tres capas: infraestructura del Tec (Azure), capa de agentes operada por HA, y las interfaces de usuario.")
add_spacer(6)
add_code("┌──────────────────────────────────────────────────────────────┐\n│                    INFRAESTRUCTURA TEC (Azure)                │\n│                                                              │\n│  ┌─────────────┐  ┌─────────────┐  ┌──────────────────────┐ │\n│  │  PostgreSQL  │  │  Vector DB  │  │  Blob Storage        │ │\n│  │             │  │  (Pinecone/ │  │                      │ │\n│  │  Perfiles    │  │   pgvector) │  │  Documentos          │ │\n│  │  Membresías  │  │             │  │  Reportes            │ │\n│  │  Conexiones  │  │  Embeddings │  │  Exportaciones       │ │\n│  │  Eventos     │  │  Relaciones │  │  Backups             │ │\n│  │  Auditoría   │  │  semánticas │  │                      │ │\n│  └──────┬──────┘  └──────┬──────┘  └──────────────────────┘ │\n│         │                │                                   │\n│         └───────┬────────┘                                   │\n│                 │                                            │\n│          ┌──────┴──────┐                                     │\n│          │   API REST  │  ← OAuth 2.0 / JWT                  │\n│          │   Gateway   │  ← Rate limiting                    │\n│          └──────┬──────┘  ← Logging completo                 │\n│                 │                                            │\n└─────────────────┼────────────────────────────────────────────┘\n                  │\n          ┌───────┴────────┐\n          │  CAPA AGENTES  │  ← Operada por HA\n          │                │\n          │  Conexión       │  Matching semántico + double opt-in\n          │  Contenidos     │  RAG + personalización por perfil\n          │  Scout          │  Análisis de redes + scoring\n          │  Insights       │  Clustering + reportes agregados\n          │                │\n          │  NO almacena   │\n          │  datos de      │\n          │  integrantes   │\n          └───────┬────────┘\n                  │\n    ┌─────────────┼──────────────┐\n    │             │              │\n    ▼             ▼              ▼\n Web App     WhatsApp Bot   Dashboard Consejo\n (React)     (Twilio/360)   (Metabase/custom)")
add_spacer(10)
add_quote("Principio clave: La capa de agentes no almacena datos personales. Procesa información anonimizada en tiempo de ejecución y devuelve resultados a la infraestructura del Tec.")

# ── 03 CAPA DE AGENTES — DETALLE ──
section_sep()
section_num("03")
add_h1("Capa de agentes — Detalle")

add_h2("Agente Conexión")
add_body("Detecta complementariedad entre integrantes basándose en conocimiento, contexto y trayectoria — no en potencial de negocio directo.")
add_spacer(6)
add_data_table(
    ["Aspecto", "Especificación"],
    [
        ["Input", "Embeddings de perfiles (conocimiento, experiencia, intereses), historial de interacciones, reglas de la comunidad"],
        ["Proceso", "Similarity semántica (cosine distance sobre embeddings) + reglas configurables por el Consejo (e.g., priorizar cross-sede, priorizar misma generación, excluir competencia directa)"],
        ["Output", "Par sugerido con score de afinidad (0-100) y explicación en lenguaje natural: \"Comparten experiencia en manufactura del Bajío y perspectivas complementarias sobre talento\""],
        ["Activación", "Double opt-in: cada parte recibe notificación anónima por WhatsApp. Solo si ambas aceptan se revela la identidad y se facilita la introducción"],
        ["Anonimización", "El LLM recibe descripciones de capacidades y necesidades sin nombres, emails ni datos de contacto"],
        ["Frecuencia", "Evaluación continua. Nuevas sugerencias se generan al actualizar perfiles o al ingresar integrantes"],
    ]
)

add_h2("Agente Contenidos")
add_body("Personaliza la información que recibe cada integrante según su perfil, industria, etapa profesional e interacciones previas.")
add_spacer(6)
add_data_table(
    ["Aspecto", "Especificación"],
    [
        ["Input", "Perfil de integrante + catálogo de contenidos (artículos, casos, recursos, grabaciones) + historial de consumo"],
        ["Proceso", "RAG (Retrieval Augmented Generation): embedding del perfil → búsqueda de contenido relevante → ranking personalizado"],
        ["Output", "Lista priorizada de contenidos con justificación: \"Dado tu sector y tu interés en nearshoring, este análisis es relevante\""],
        ["Diferenciación", "Una persona con 3 años de experiencia recibe material distinto de quien lleva 15"],
        ["Aprendizaje", "El agente mejora con cada interacción: qué se abrió, qué se ignoró, qué se compartió"],
    ]
)

add_h2("Agente Scout")
add_body("Identifica personas candidatas para futuras generaciones a partir de las redes de integrantes actuales.")
add_spacer(6)
add_data_table(
    ["Aspecto", "Especificación"],
    [
        ["Input", "Redes públicas de integrantes (LinkedIn), criterios de admisión definidos por el Consejo"],
        ["Proceso", "Análisis de red (graph analysis) + matching semántico con perfil Tec Beyond"],
        ["Output", "Lista de candidaturas con contexto: \"En la red de [Integrante X] hay alguien con perfil afín — 85% de afinidad\""],
        ["Gobernanza", "El Consejo decide a quién invitar. El agente nunca contacta directamente"],
        ["Fuentes", "Solo datos públicos. No scraping de datos privados"],
    ]
)

add_h2("Agente Insights")
add_body("Genera inteligencia estratégica para el Consejo a partir de datos agregados y anonimizados.")
add_spacer(6)
add_data_table(
    ["Aspecto", "Especificación"],
    [
        ["Input", "Datos agregados de toda la red: perfiles, interacciones, contenido consumido, conexiones formadas"],
        ["Proceso", "Clustering, análisis de tendencias, detección de patrones, NLP sobre interacciones"],
        ["Output", "Reportes periódicos: \"47 integrantes mencionaron nearshoring este trimestre\", \"La sede Chihuahua tiene 30% menos participación — considerar acercamiento\""],
        ["Frecuencia", "Reportes mensuales automáticos + alertas en tiempo real para anomalías"],
        ["Privacidad", "Trabaja exclusivamente con datos agregados. Nunca genera reportes sobre una persona individual"],
    ]
)

# ── 04 STACK TECNOLÓGICO ──
section_sep()
section_num("04")
add_h1("Stack tecnológico")
add_spacer(6)
add_data_table(
    ["Capa", "Tecnología", "Justificación"],
    [
        ["Backend", "Node.js / Python (FastAPI)", "Ecosistema maduro, compatible con Azure"],
        ["BD relacional", "PostgreSQL (Azure Database)", "Estándar del Tec, robusto, escalable"],
        ["BD vectorial", "pgvector o Pinecone", "Búsqueda semántica para conexiones y contenido"],
        ["LLM", "Claude (Anthropic) / GPT-4o (OpenAI)", "Procesamiento de lenguaje natural para agentes"],
        ["Orquestación agentes", "Horizons Architecture framework", "Coordinación multi-agente con reglas de negocio"],
        ["Frontend web", "React + Tailwind CSS", "Responsive, rendimiento, accesibilidad"],
        ["WhatsApp", "WhatsApp Business API (Twilio/360dialog)", "Canal principal de la comunidad"],
        ["Dashboard", "Metabase o custom (React + Chart.js)", "Visualización de insights para Consejo"],
        ["Autenticación", "OAuth 2.0 + JWT", "SSO compatible con ecosistema Tec (Azure AD)"],
        ["Infraestructura", "Microsoft Azure", "Alineado con infraestructura existente del Tec"],
        ["CI/CD", "GitHub Actions → Azure App Service", "Despliegue automatizado"],
        ["Monitoreo", "Azure Monitor + Application Insights", "Observabilidad y alertas"],
    ]
)

# ── 05 MODELO DE DATOS ──
section_sep()
section_num("05")
add_h1("Modelo de datos (entidades principales)")
add_spacer(6)
add_code("INTEGRANTE\n├── id, nombre, email, teléfono\n├── empresa, cargo, sector, ubicación\n├── sede_id, generación, año_ingreso\n├── conocimiento_experiencia (texto libre → embedding)\n├── que_quiere_aprender (texto libre → embedding)\n├── preferencias_privacidad (granular: visible / solo sistema / privado)\n├── canal_preferido (whatsapp / web / ambos)\n├── estado (activo / inactivo / pendiente)\n└── fecha_registro, última_actividad\n\nLEGADO_INDIVIDUAL\n├── id, integrante_id\n├── metas_personales (texto libre)\n├── contribuciones [] (conexiones formadas, conocimiento compartido, mentorías)\n├── trayectoria [] (hitos registrados a lo largo del tiempo)\n└── fecha_creación, última_actualización\n\nLEGADO_GENERACIONAL\n├── id, generación, año_inicio\n├── integrantes_count, sedes []\n├── temas_clave [] (detectados por Insights)\n├── conexiones_formadas_count\n├── resumen_narrativo (generado por agente, revisado por Consejo)\n└── última_actualización\n\nSEDE\n├── id, nombre, campus, ciudad\n├── admin_id\n├── fecha_inicio\n└── estado (activa / en preparación)\n\nCONEXION\n├── id, integrante_a_id, integrante_b_id\n├── score_afinidad (0-100)\n├── base_conexion (texto: por qué se sugiere)\n├── estado_a (pendiente / aceptado / rechazado)\n├── estado_b (pendiente / aceptado / rechazado)\n├── estado_final (facilitado / cancelado / archivado)\n├── facilitado_por (admin_id)\n└── fechas (creado, respondido_a, respondido_b, facilitado)\n\nCONTENIDO\n├── id, título, tipo (artículo / caso / recurso / video)\n├── sectores_relevantes [], nivel_experiencia\n├── embedding (para matching con perfiles)\n└── fecha_publicación, fuente\n\nCANDIDATURA_SCOUT\n├── id, nombre_estimado, empresa, sector\n├── detectado_via_integrante_id\n├── score_afinidad (0-100)\n├── estado (detectado / revisado / invitado / descartado)\n└── decidido_por (consejo), fecha_decisión\n\nEVENTO\n├── id, nombre, tipo (presencial / virtual), sede_id\n├── fecha, lugar, capacidad\n└── asistentes [] (integrante_id + confirmación)\n\nLOG_AUDITORÍA\n├── id, persona_id, acción, recurso\n├── timestamp, ip, resultado\n└── detalle (JSON)")

# ── 06 INTEGRACIONES ──
section_sep()
section_num("06")
add_h1("Integraciones")
add_spacer(6)
add_data_table(
    ["Sistema", "Tipo", "Qué se intercambia", "Dirección"],
    [
        ["Azure AD (Tec)", "SSO", "Autenticación de cuentas institucionales", "Tec → Plataforma"],
        ["WhatsApp Business API", "Mensajería", "Notificaciones, respuestas, flujos conversacionales", "Bidireccional"],
        ["LinkedIn (opcional)", "Enriquecimiento", "Datos públicos de perfil profesional", "LinkedIn → Plataforma"],
        ["Anthropic / OpenAI API", "LLM", "Procesamiento de lenguaje para agentes", "Plataforma → API"],
        ["SMTP / SendGrid", "Email", "Notificaciones, reportes, invitaciones", "Plataforma → Integrante"],
        ["Azure Blob Storage", "Archivos", "Documentos, reportes exportados, backups", "Interno"],
    ]
)

# ── 07 SEGURIDAD Y PRIVACIDAD ──
section_sep()
section_num("07")
add_h1("Seguridad y privacidad")
add_spacer(4)
add_quote(
    "\"Mucha gente no respondió la encuesta cuando leyó los avisos de privacidad del Tec. Tiene una vulnerabilidad reconocida en la comunidad.\"",
    "— Rodrigo, Consejo de representantes"
)
add_spacer(10)
add_quote("Principio de diseño: Cada integrante controla su información. Decide qué comparte, con quién, y puede retirarlo en cualquier momento. Los agentes de IA trabajan con datos anonimizados — nunca ven nombres, emails ni datos de contacto.")
add_spacer(10)
add_data_table(
    ["Capa", "Medida", "Estándar"],
    [
        ["Datos en reposo", "AES-256", "Azure Storage Service Encryption"],
        ["Datos en tránsito", "TLS 1.3", "HTTPS obligatorio"],
        ["Anonimización para IA", "Datos procesados sin PII antes de enviar a LLM", "Política interna verificable"],
        ["Autenticación admin", "MFA (Multi-Factor Authentication)", "Azure AD + Authenticator"],
        ["API", "JWT con rotación de tokens", "Expiración configurable"],
        ["Acceso a datos", "RBAC (Role-Based Access Control)", "6 roles definidos"],
        ["Logs", "Registro inmutable de toda operación", "Azure Monitor + tabla interna"],
        ["Backups", "Automático diario, retención 30 días", "Azure Backup"],
        ["Pruebas de seguridad", "Pentest antes de lanzamiento", "Proveedor externo"],
        ["Transparencia", "Cada integrante puede ver y exportar toda su información", "Panel de privacidad en perfil"],
    ]
)

add_h2("Roles y permisos")
add_spacer(6)
add_data_table(
    ["Rol", "Datos integrantes", "Conexiones", "Insights", "Config", "Auditoría"],
    [
        ["Super Admin", "Todos", "Todas", "Todos", "Sí", "Sí"],
        ["Admin Consejo", "Agregados", "Aprobar", "Todos", "Reglas", "Sí"],
        ["Admin Sede", "Solo su sede", "Solo su sede", "Solo su sede", "No", "Solo su sede"],
        ["Operación HA", "No", "No", "No", "Técnica", "Técnica"],
        ["Integrante", "Solo propio", "Propias", "No", "Su perfil", "No"],
        ["Auditoría", "Lectura", "Lectura", "Lectura", "No", "Sí"],
    ]
)

# ── 08 RESIDENCIA Y PROPIEDAD DE DATOS ──
section_sep()
section_num("08")
add_h1("Residencia y propiedad de datos")
add_spacer(6)
add_data_table(
    ["Componente", "Ubicación", "Propiedad", "Si termina el contrato"],
    [
        ["Datos de integrantes", "Azure Tec", "Tec Beyond 100%", "Se revoca acceso a HA; datos intactos"],
        ["Base vectorial (embeddings)", "Azure Tec", "Tec Beyond 100%", "Se entrega documentación para migración"],
        ["Insights y reportes", "Azure Tec", "Tec Beyond 100%", "Quedan en infraestructura Tec"],
        ["Código de plataforma", "Repositorio HA", "Horizons Architecture", "Licencia se termina"],
        ["Algoritmos de agentes", "Repositorio HA", "Horizons Architecture", "HA retira su código; resultados quedan"],
        ["Reglas de comunidad", "Documentadas en plataforma", "Consejo León", "Exportables en formato estándar"],
    ]
)
add_spacer(10)
add_quote("Compromiso de portabilidad: Si Tec Beyond migra a otro proveedor, HA entrega todos los datos en formato estándar (CSV/JSON), documentación técnica de la base vectorial, y un periodo de transición acordado en contrato.")

# ── 09 ESCALABILIDAD ──
section_sep()
section_num("09")
add_h1("Escalabilidad")

add_h3("Piloto (León)")
add_data_table(
    ["Recurso", "Dimensionamiento"],
    [
        ["Conexiones simultáneas", "~50-100"],
        ["Integrantes totales", "~300"],
        ["Embeddings", "~300 perfiles + contenido"],
        ["Almacenamiento", "< 10 GB"],
        ["Costo mensual estimado Azure", "~$200-400 USD"],
        ["Costo LLM mensual estimado", "~$100-300 USD (depende de volumen)"],
    ]
)

add_h3("Expansión (3-5 sedes)")
add_data_table(
    ["Recurso", "Dimensionamiento"],
    [
        ["Conexiones simultáneas", "~200-500"],
        ["Integrantes totales", "~1,500"],
        ["Costo marginal por sede nueva", "~20-30% del costo del piloto"],
    ]
)
add_spacer(6)
add_body("La arquitectura es horizontal: agregar una sede es crear un nuevo nodo con su administración, no redesplegar el sistema. La base de conocimiento es compartida — cada nueva persona enriquece las conexiones de toda la red.")

# ── 10 FASES DE DESARROLLO ──
section_sep()
section_num("10")
add_h1("Fases de desarrollo")
add_spacer(6)
add_data_table(
    ["Fase", "Duración", "Entregables"],
    [
        ["Descubrimiento", "2-3 semanas", "Diseño detallado, modelo de datos final, cotización cerrada"],
        ["MVP", "8-10 semanas", "Perfiles, legado personal, conexiones básicas, WhatsApp bot, dashboard Consejo"],
        ["Agentes v1", "4-6 semanas", "Conexiones + Contenidos operando con base de conocimiento real"],
        ["Agentes v2", "4-6 semanas", "Scout + Insights operando"],
        ["Estabilización", "2-4 semanas", "Testing, pentest, ajustes de performance, onboarding"],
        ["Lanzamiento piloto", "—", "~300 integrantes de León participando"],
    ]
)
add_spacer(6)
add_quote("Total estimado: 5-8 meses desde aprobación hasta lanzamiento del piloto.")

add_h3("Post-lanzamiento: plataforma viva")
add_bullet("Actualizaciones mensuales de algoritmos de conexión y contenido basadas en patrones reales de uso")
add_bullet("Reportes trimestrales al Consejo con insights y recomendaciones de evolución")
add_bullet("Nuevas funcionalidades según las necesidades que emerjan de la comunidad")
add_bullet("Expansión a nuevas sedes como nodos adicionales", last=True)

# ── 11 PREGUNTAS FRECUENTES PARA TI ──
section_sep()
section_num("11")
add_h1("Preguntas frecuentes para TI")

add_h3("¿Se necesita infraestructura nueva en el Tec?")
add_body("No. Todo corre sobre Azure, que el Tec ya tiene. Se crea un resource group dedicado para Tec Beyond.")

add_h3("¿Los agentes de IA corren dentro de Azure?")
add_body("Los agentes se despliegan en Azure App Service dentro del tenant del Tec. Las llamadas a LLMs (Claude/GPT-4o) salen a APIs externas, pero los datos que se envían son procesados y anonimizados — nunca se envía un perfil completo a un LLM externo.")

add_h3("¿Qué pasa con la latencia de WhatsApp?")
add_body("El bot responde en < 5 segundos para interacciones simples (aceptar conexión, confirmar evento). Para respuestas que requieren procesamiento de agentes, < 15 segundos.")

add_h3("¿Cómo se actualiza la base vectorial?")
add_body("Los embeddings se regeneran cuando alguien actualiza su perfil. El sistema también re-indexa periódicamente (semanal) para incorporar datos de interacciones recientes.")

add_h3("¿Qué datos se envían a las APIs de LLM?")
add_body("Solo contexto procesado y anonimizado. Para generar una conexión, el agente envía al LLM las descripciones de capacidades y necesidades sin nombres, emails ni datos de contacto. El LLM nunca ve datos personales identificables.")

add_h3("¿Se puede integrar con sistemas existentes del Tec (CRM, ERP)?")
add_body("Sí, vía API REST. La fase de descubrimiento mapea qué integraciones específicas se necesitan.")

# Colophon
add_spacer(24)
p = doc.add_paragraph()
pPr = p._element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/></w:pBdr>')
pPr.append(pBdr)
add_spacer(10)
add_body("Documento preparado por Horizons Architecture para Tec Beyond.")
mono_line("Edgar Barroso — edgar@horizonsarchitecture.ai")
mono_line("Marzo 2026 — horizonsarchitecture.ai")

# ── SAVE ──
doc.save(out_path)
print(f"Generated: {out_path}")
