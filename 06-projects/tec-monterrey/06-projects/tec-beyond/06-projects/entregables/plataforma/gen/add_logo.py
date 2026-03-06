"""Post-process: replace logo placeholder with actual image using python-docx."""
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import os, sys

script_dir = os.path.dirname(os.path.abspath(__file__))
docx_path = os.path.join(script_dir, "..", "Plataforma-Tec-Beyond-Descripcion-v1.docx")
logo_path = os.path.join(script_dir, "logo.png")

doc = Document(docx_path)

# Find and replace the placeholder paragraph with the logo
for i, para in enumerate(doc.paragraphs):
    if "{{LOGO_PLACEHOLDER}}" in para.text:
        # Clear the placeholder text
        for run in para.runs:
            run.text = ""
        # Add image run
        run = para.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
        print(f"Logo inserted at paragraph {i}")
        break
else:
    print("Placeholder not found, inserting at beginning")
    # Fallback: insert at start
    body = doc.element.body
    new_para = doc.add_paragraph()
    run = new_para.add_run()
    run.add_picture(logo_path, width=Inches(0.85))
    body.insert(0, new_para._element)

doc.save(docx_path)
print(f"Saved: {docx_path}")
